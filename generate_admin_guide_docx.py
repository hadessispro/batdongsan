from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT_DIR = ROOT / "docs"
OUT_FILE = OUT_DIR / "Huong-dan-su-dung-admin-Bich-Dong-Lakeside.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_bullet(document: Document, text: str, level: int = 0, bold_prefix: str | None = None) -> None:
    p = document.add_paragraph(style="List Bullet")
    if level:
        p.paragraph_format.left_indent = Inches(0.25 * level)
    if bold_prefix and text.startswith(bold_prefix):
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)


def add_number(document: Document, text: str, level: int = 0) -> None:
    p = document.add_paragraph(style="List Number")
    if level:
        p.paragraph_format.left_indent = Inches(0.25 * level)
    p.add_run(text)


def add_note(document: Document, label: str, text: str, color: str = "EAF6F5") -> None:
    table = document.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, color)
    p = cell.paragraphs[0]
    r1 = p.add_run(label + ": ")
    r1.bold = True
    p.add_run(text)
    document.add_paragraph()


def add_title_page(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("HƯỚNG DẪN SỬ DỤNG WEBSITE ADMIN")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x0D, 0x94, 0x88)

    p2 = document.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p2.add_run("Bích Động Lakeside")
    run.bold = True
    run.font.size = Pt(18)

    p3 = document.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p3.add_run(
        "Tài liệu vận hành dành cho người dùng không rành kỹ thuật\n"
        "Phiên bản tài liệu: "
        + datetime.now().strftime("%d/%m/%Y %H:%M")
    )
    run.font.size = Pt(11)

    add_note(
        document,
        "Mục tiêu tài liệu",
        "Giúp người vận hành có thể đăng nhập, cập nhật nội dung, upload media, quản lý thư viện ảnh/video, "
        "chỉnh VR 360, tham quan 360, master plan và JSON nâng cao mà không cần biết lập trình.",
    )


def add_overview_tables(document: Document) -> None:
    document.add_heading("1. Tổng Quan Hệ Thống", level=1)
    document.add_paragraph(
        "Khu vực quản trị của website đã được tách riêng bằng trang đăng nhập và phiên đăng nhập session. "
        "Người dùng thông thường không nên truy cập trực tiếp vào các file hệ thống."
    )

    t = document.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    t.rows[0].cells[0].text = "Hạng mục"
    t.rows[0].cells[1].text = "Thông tin"
    rows = [
        ("URL đăng nhập", "/admin"),
        ("URL panel sau khi đăng nhập", "/admin/panel"),
        ("Đăng xuất", "Nút Đăng xuất ở góc trên bên phải của admin"),
        ("Cơ chế đăng nhập", "Session server-side qua API"),
        ("Nơi lưu tài khoản/mật khẩu", "api/config.local.php trên hosting"),
        ("Nơi xử lý API admin", "api/index.php"),
    ]
    for left, right in rows:
        row = t.add_row().cells
        row[0].text = left
        row[1].text = right
    document.add_paragraph()

    add_note(
        document,
        "Quan trọng",
        "Không đưa file api/config.local.php lên GitHub công khai. File này nên chỉ tồn tại trên hosting hoặc máy quản trị nội bộ.",
        color="FFF4E5",
    )


def add_login_section(document: Document) -> None:
    document.add_heading("2. Cách Đăng Nhập Và Đăng Xuất", level=1)
    add_number(document, "Mở trình duyệt và vào đường dẫn /admin trên domain của website.")
    add_number(document, "Nhập tài khoản và mật khẩu admin đã được quản trị viên cấp.")
    add_number(document, "Bấm Đăng nhập vào admin.")
    add_number(document, "Khi làm xong, bấm nút Đăng xuất ở góc trên bên phải.")

    add_bullet(document, "Người dùng không cần nhớ /admin-login.php hay /admin.php. Chỉ cần nhớ /admin.")
    add_bullet(document, "Nếu đang đăng nhập rồi mà mở lại /admin, hệ thống sẽ tự chuyển thẳng vào panel.")
    add_bullet(document, "Nếu quên mật khẩu, cần nhờ quản trị viên kỹ thuật đổi lại trong file api/config.local.php.")

    add_note(
        document,
        "Khuyến nghị bảo mật",
        "Không lưu mật khẩu trong file chat, không gửi qua nhóm đông người, và không dùng lại mật khẩu của email hay hosting.",
        color="FDECEC",
    )


def add_menu_map(document: Document) -> None:
    document.add_heading("3. Sơ Đồ Menu Admin", level=1)
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Menu"
    hdr[1].text = "Mục đích"
    hdr[2].text = "Đối tượng nên dùng"

    items = [
        ("Dashboard", "Xem tổng quan nhanh và lối tắt sang các khu vực khác.", "Mọi người"),
        ("Quản lý File", "Upload, xem trước, lọc và xóa media trong thư mục frames.", "Mọi người"),
        ("Album / Thư Viện", "Quản lý album ảnh và video thư viện dùng trên website.", "Mọi người"),
        ("VR 360° Config", "Chỉnh lon/lat/FOV của các tiện ích 360.", "Người biết nghiệp vụ VR"),
        ("Hotspot Editor", "Chỉnh JSON hotspot của VR 360.", "Người nâng cao"),
        ("Tham quan 360", "Quản lý pano tham quan, hotspot và chữ 3D.", "Mọi người; ưu tiên dùng tool trực quan"),
        ("Master Plan", "Chỉnh vị trí top/left, icon, group của các tiện ích trên bản đồ tổng quan.", "Mọi người"),
        ("Cài đặt chung", "Hiện là màn hình mẫu/tham khảo, chưa nối lưu thực tế.", "Chỉ xem, không dùng để lưu"),
        ("JSON Editor", "Chỉnh thẳng file JSON nâng cao.", "Người hiểu dữ liệu, dùng cực thận trọng"),
    ]

    for menu, purpose, audience in items:
        row = table.add_row().cells
        row[0].text = menu
        row[1].text = purpose
        row[2].text = audience
    document.add_paragraph()


def add_file_manager(document: Document) -> None:
    document.add_heading("4. Hướng Dẫn Quản Lý File", level=1)
    document.add_paragraph(
        "Khu vực Quản lý File là nơi upload và quản lý media trong thư mục frames. Đây là màn hình quan trọng nhất "
        "khi cần đưa ảnh, video hoặc PDF mới lên website."
    )

    add_bullet(document, "Thanh tìm kiếm: tìm nhanh theo tên file.")
    add_bullet(document, "Bộ lọc: lọc theo Ảnh, Video, PDF hoặc Thư mục.")
    add_bullet(document, "Upload media: upload file mới vào thư mục đang mở.")
    add_bullet(document, "Preview: bấm vào file để xem trước ảnh, video hoặc PDF.")
    add_bullet(document, "Delete: bấm biểu tượng thùng rác để xóa file media.")

    document.add_paragraph("Các loại file hỗ trợ upload qua admin:")
    for ext in ["jpg, jpeg, png, webp, gif, svg", "mp4, webm, mov", "pdf"]:
        add_bullet(document, ext)

    document.add_paragraph("Các thư mục media quan trọng cần nhớ:")
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    h = table.rows[0].cells
    h[0].text = "Thư mục"
    h[1].text = "Dùng cho"
    h[2].text = "Gợi ý sử dụng"
    folders = [
        ("frames/galley", "Thư viện ảnh và video", "Hầu hết ảnh album và video thư viện nên nằm ở đây"),
        ("frames/3dvr", "Ảnh pano VR 360", "Không nên đổi tên lung tung nếu không hiểu luồng VR"),
        ("frames/tienich", "Pano tham quan 360", "Dùng cho màn tham quan"),
        ("frames/thamquan", "Tool tham quan 360", "Có thể chứa file tool/hỗ trợ"),
        ("frames/matbang", "Ảnh và dữ liệu mặt bằng", "Liên quan masterplan và lot-data"),
        ("frames/phankhu", "Ảnh và SVG phân khu", "Dùng cho màn phân khu"),
        ("frames/brochur", "Brochure PDF", "Tài liệu PDF cho người xem tải về"),
    ]
    for folder, use, note in folders:
        row = table.add_row().cells
        row[0].text = folder
        row[1].text = use
        row[2].text = note
    document.add_paragraph()

    add_note(
        document,
        "Lưu ý quan trọng",
        "Quản lý File chỉ làm việc với vùng frames. Các file dữ liệu JSON trong data không được upload từ đây; "
        "muốn sửa JSON thì dùng JSON Editor.",
    )


def add_library_section(document: Document) -> None:
    document.add_heading("5. Quản Lý Album Ảnh Và Video Thư Viện", level=1)
    document.add_paragraph(
        "Đây là nơi quản lý nội dung xuất hiện ở màn Thư viện trên website. Khu vực này gồm 2 phần: album ảnh và video thư viện."
    )

    document.add_heading("5.1. Album Ảnh", level=2)
    add_bullet(document, "Thêm Album: tạo album mới.")
    add_bullet(document, "Nút dấu cộng trên từng album: thêm ảnh bằng cách nhập đường dẫn file đã có sẵn.")
    add_bullet(document, "Nút upload trên từng album: upload ảnh mới rồi gắn luôn vào album.")
    add_bullet(document, "Nút bút: đổi tên album.")
    add_bullet(document, "Nút thùng rác: xóa album.")
    add_bullet(document, "Nút dấu x trên thumbnail: gỡ ảnh khỏi album.")

    add_note(
        document,
        "Quy tắc lưu album",
        "Sau khi thêm/sửa/xóa album hoặc ảnh trong album, nên bấm Lưu thư viện để ghi dữ liệu vào data/library.json. "
        "Nếu chỉ thao tác trên giao diện mà chưa lưu, dữ liệu có thể chưa cập nhật chính thức.",
        color="FFF4E5",
    )

    document.add_heading("5.2. Video Thư Viện", level=2)
    add_bullet(document, "Upload & thêm video: cách dễ nhất cho người ít kỹ thuật. Chọn file là hệ thống tự upload, tự tạo video trong thư viện và tự lưu.")
    add_bullet(document, "Chọn video đã upload: dùng khi file video đã tồn tại sẵn trong frames/galley.")
    add_bullet(document, "Tiêu đề video: sửa trực tiếp trong bảng.")
    add_bullet(document, "Thumbnail: chọn từ danh sách ảnh thư viện có sẵn, có preview nhỏ đi kèm.")
    add_bullet(document, "Xóa video: bấm thùng rác ở dòng video.")

    add_note(
        document,
        "Điểm rất quan trọng",
        "Người dùng không cần nhớ đường dẫn video hay gõ thumbnail thủ công nữa. "
        "Cách làm an toàn nhất là: bấm Upload & thêm video, chờ toast báo thành công, rồi kiểm tra lại tiêu đề/thumbnail.",
    )

    add_note(
        document,
        "Lưu ý tên thư mục",
        "Thư mục đúng của thư viện là frames/galley. Đây là cách viết đang dùng trong hệ thống hiện tại. "
        "Không tự đổi thành gallery nếu chưa có điều chỉnh code tương ứng.",
        color="FDECEC",
    )


def add_vr_sections(document: Document) -> None:
    document.add_heading("6. VR 360° Config", level=1)
    document.add_paragraph(
        "Màn hình này phục vụ cấu hình góc nhìn khi nhảy vào tiện ích 360. Mỗi tiện ích có thể có bộ số lon/lat/FOV riêng."
    )
    add_bullet(document, "Mở Visual Editor: cách nên dùng đầu tiên.")
    add_bullet(document, "Bảng lon/lat/FOV nâng cao: chỉ dùng khi cần nhập số chính xác.")
    add_bullet(document, "Lưu thay đổi: ghi vào dữ liệu VR hotspot chung.")
    add_note(
        document,
        "Khuyên dùng",
        "Nếu không chắc lon/lat/FOV là gì, hãy dùng Visual Editor trước. "
        "Bảng số chỉ nên dùng cho người đã hiểu góc nhìn VR.",
    )

    document.add_heading("7. Hotspot Editor (JSON nâng cao)", level=1)
    document.add_paragraph(
        "Màn này hiển thị toàn bộ dữ liệu từ vr-hotspot-data.json. Đây là chế độ nâng cao để sửa trực tiếp JSON."
    )
    add_bullet(document, "Dùng khi cần chỉnh hotspot theo dạng dữ liệu thô.")
    add_bullet(document, "Có thể bấm Copy để gửi dữ liệu cho người kỹ thuật.")
    add_bullet(document, "Sau khi chỉnh xong phải bấm Lưu JSON.")
    add_note(
        document,
        "Rủi ro",
        "Sai một dấu ngoặc hoặc dấu phẩy cũng có thể làm hỏng JSON. Nếu không rành, hãy dùng Visual Editor thay vì JSON Editor thô.",
        color="FDECEC",
    )

    document.add_heading("8. Tham Quan 360", level=1)
    document.add_paragraph(
        "Màn này có 2 cách làm việc: tool trực quan nhúng ngay trong admin và form nâng cao để sửa dữ liệu dự phòng."
    )
    add_bullet(document, "Cách nên dùng cho người low-tech: thao tác trực tiếp trong tool 360 rồi bấm lưu.")
    add_bullet(document, "Ảnh pano tham quan: đường dẫn file pano đang dùng.")
    add_bullet(document, "Hotspot PX/PY: tọa độ pixel trên ảnh pano chuẩn 2000×1000.")
    add_bullet(document, "Chữ 3D JSON: vùng JSON nâng cao cho text nổi.")
    add_bullet(document, "Lưu form nâng cao: ghi về data/thamquan_config.json.")
    add_note(
        document,
        "Gợi ý vận hành",
        "Nếu chỉ cần chỉnh vài điểm hotspot, hãy dùng tool trực quan. "
        "Chỉ mở phần PX/PY và JSON khi cần sửa kỹ hoặc khi được người kỹ thuật hướng dẫn.",
    )


def add_masterplan_and_settings(document: Document) -> None:
    document.add_heading("9. Master Plan", level=1)
    document.add_paragraph(
        "Master Plan dùng để đặt vị trí các tiện ích trên sơ đồ tổng quan bằng tỷ lệ phần trăm."
    )
    add_bullet(document, "Top %: vị trí theo chiều dọc.")
    add_bullet(document, "Left %: vị trí theo chiều ngang.")
    add_bullet(document, "Nhóm: mã nhóm hiển thị, ví dụ g1.")
    add_bullet(document, "Icon: class icon Font Awesome.")
    add_bullet(document, "Lưu form: cách dùng dễ nhất cho người vận hành.")
    add_bullet(document, "Lưu JSON: dùng khi chỉnh nâng cao.")
    add_note(
        document,
        "Mẹo thực tế",
        "Sau khi thay đổi vị trí top/left, nên kiểm tra ngay ngoài website để đảm bảo icon không bị lệch hoặc chồng lên nhau.",
    )

    document.add_heading("10. Cài Đặt Chung", level=1)
    document.add_paragraph(
        "Màn này hiện có giao diện nhập liệu cho Tên dự án, Hotline và Favicon, nhưng ở phiên bản hiện tại chưa có luồng save thực tế nối với backend."
    )
    add_note(
        document,
        "Trạng thái hiện tại",
        "Hãy xem Cài đặt chung như phần giao diện mẫu/tham khảo. "
        "Không nên coi đây là nơi lưu chính thức cho các thông tin quan trọng cho đến khi được nối backend.",
        color="FFF4E5",
    )


def add_json_editor(document: Document) -> None:
    document.add_heading("11. JSON Editor", level=1)
    document.add_paragraph(
        "JSON Editor là công cụ nâng cao để sửa trực tiếp file dữ liệu. Chỉ nên dùng khi người thao tác hiểu rõ dữ liệu đang sửa."
    )

    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "File"
    hdr[1].text = "Ý nghĩa"

    files = [
        ("vr-hotspot-data.json", "Hotspot VR 360 và cấu hình related VR"),
        ("data/vr_config.json", "Cấu hình VR bổ sung"),
        ("data/masterplan.json", "Dữ liệu vị trí tiện ích trên master plan"),
        ("data/buildings.json", "Dữ liệu tòa/khối nhà"),
        ("data/library.json", "Album ảnh và video thư viện"),
        ("data/vitri_config.json", "Cấu hình màn vị trí"),
        ("data/thamquan_config.json", "Pano tham quan, hotspot và chữ 3D"),
        ("frames/matbang/lot-data.json", "Dữ liệu lô trong mặt bằng"),
    ]
    for file_name, desc in files:
        row = table.add_row().cells
        row[0].text = file_name
        row[1].text = desc
    document.add_paragraph()

    add_bullet(document, "Validate JSON: kiểm tra cú pháp.")
    add_bullet(document, "Format: sắp xếp lại JSON cho dễ đọc.")
    add_bullet(document, "Lưu file: ghi file JSON đã chọn về server.")
    add_bullet(document, "Copy: sao chép toàn bộ JSON ra clipboard.")

    add_note(
        document,
        "Nguyên tắc an toàn",
        "Nếu không chắc file nào đang chỉnh, đừng bấm Lưu. "
        "Sai file JSON có thể làm lỗi toàn bộ một khu vực của website.",
        color="FDECEC",
    )


def add_workflows(document: Document) -> None:
    document.add_heading("12. Các Tác Vụ Thường Gặp", level=1)

    document.add_heading("12.1. Thêm Ảnh Mới Vào Album", level=2)
    add_number(document, "Vào Album / Thư Viện.")
    add_number(document, "Chọn album cần thêm ảnh.")
    add_number(document, "Bấm nút upload ở album đó để tải ảnh mới lên, hoặc bấm dấu cộng nếu ảnh đã có sẵn.")
    add_number(document, "Kiểm tra thumbnail đã hiện đúng.")
    add_number(document, "Bấm Lưu thư viện.")

    document.add_heading("12.2. Thêm Video Mới Cho Thư Viện", level=2)
    add_number(document, "Vào Album / Thư Viện.")
    add_number(document, "Bấm Upload & thêm video.")
    add_number(document, "Chọn file mp4/webm/mov từ máy tính.")
    add_number(document, "Chờ thông báo thành công.")
    add_number(document, "Kiểm tra lại tiêu đề và thumbnail.")
    add_number(document, "Nếu cần, đổi thumbnail bằng dropdown ảnh.")

    document.add_heading("12.3. Sửa Hotspot Tham Quan 360", level=2)
    add_number(document, "Vào Tham quan 360.")
    add_number(document, "Ưu tiên thao tác trên tool 360 trực quan.")
    add_number(document, "Chỉ dùng form nâng cao nếu cần sửa chính xác hoặc sửa thủ công.")
    add_number(document, "Bấm Lưu form nâng cao khi dùng phần dự phòng.")

    document.add_heading("12.4. Sửa Vị Trí Tiện Ích Trên Master Plan", level=2)
    add_number(document, "Vào Master Plan.")
    add_number(document, "Tìm đúng dòng tiện ích.")
    add_number(document, "Điều chỉnh Top % và Left %.")
    add_number(document, "Bấm Lưu form.")
    add_number(document, "Mở website kiểm tra vị trí hiển thị.")


def add_troubleshooting(document: Document) -> None:
    document.add_heading("13. Xử Lý Sự Cố Thường Gặp", level=1)

    issues = [
        ("Không đăng nhập được", "Kiểm tra lại tài khoản/mật khẩu. Nếu vẫn sai, quản trị viên kỹ thuật cần xem file api/config.local.php trên hosting."),
        ("Đăng nhập thành công nhưng không vào được panel", "Kiểm tra rewrite Apache/.htaccess và chắc chắn hosting cho phép PHP session hoạt động."),
        ("Upload xong nhưng không thấy file", "Kiểm tra đang đứng đúng thư mục frames tương ứng. Sau upload nên đợi toast báo thành công rồi refresh màn hình."),
        ("Video không lên ngoài website", "Kiểm tra video đã có trong data/library.json chưa. Với người low-tech, luôn dùng Upload & thêm video để tránh quên bước lưu."),
        ("Không lưu được JSON", "Thường do quyền ghi thư mục data hoặc data/backups trên hosting."),
        ("Bấm Lưu nhưng website chưa đổi", "Thử Ctrl+F5, kiểm tra đúng file dữ liệu đã lưu, và xác nhận không sửa nhầm môi trường/staging."),
        ("Không đăng xuất được", "Phiên bản mới đã có nút Đăng xuất hoạt động. Nếu vẫn lỗi, làm mới file admin.php trên hosting và thử lại."),
    ]

    for title, answer in issues:
        p = document.add_paragraph()
        r = p.add_run(title + ": ")
        r.bold = True
        p.add_run(answer)


def add_security_notes(document: Document) -> None:
    document.add_heading("14. Quy Tắc Bảo Mật Khi Vận Hành", level=1)
    for item in [
        "Chỉ truy cập khu vực quản trị qua /admin.",
        "Không gửi mật khẩu qua nhóm chat đông người.",
        "Không lưu tài khoản/mật khẩu trong file Word gửi rộng rãi.",
        "Sau khi dùng xong nên bấm Đăng xuất, đặc biệt trên máy dùng chung.",
        "Không sửa JSON nâng cao nếu chưa hiểu dữ liệu.",
        "Không xóa file media cũ nếu chưa chắc website không còn dùng file đó.",
        "Đổi mật khẩu định kỳ nếu website có nhiều người cùng dùng admin.",
    ]:
        add_bullet(document, item)


def add_appendix(document: Document) -> None:
    document.add_heading("15. Phụ Lục: Bản Đồ Dữ Liệu", level=1)
    document.add_paragraph("Các nơi lưu chính trong hệ thống:")

    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    h = table.rows[0].cells
    h[0].text = "Đường dẫn"
    h[1].text = "Nội dung"

    rows = [
        ("api/config.local.php", "Tài khoản và mật khẩu admin; không đưa public"),
        ("api/index.php", "API lưu/đọc dữ liệu admin"),
        ("data/library.json", "Album ảnh và video thư viện"),
        ("data/masterplan.json", "Master plan"),
        ("data/thamquan_config.json", "Tham quan 360"),
        ("data/vitri_config.json", "Cấu hình vị trí"),
        ("vr-hotspot-data.json", "Hotspot VR 360"),
        ("frames/galley", "Ảnh và video thư viện"),
        ("frames/3dvr", "Pano VR 360"),
        ("frames/tienich", "Pano tham quan"),
    ]
    for path, desc in rows:
        row = table.add_row().cells
        row[0].text = path
        row[1].text = desc
    document.add_paragraph()

    add_note(
        document,
        "Khuyến nghị bàn giao",
        "Khi bàn giao cho người dùng low-tech, nên gửi kèm tài liệu này và tổ chức 1 buổi hướng dẫn ngắn "
        "theo 4 thao tác chính: đăng nhập, upload ảnh, thêm video, lưu thư viện.",
    )


def build_document() -> Path:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].font.color.rgb = RGBColor(0x0D, 0x47, 0x43)
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(12.5)
    styles["Heading 2"].font.bold = True

    add_title_page(doc)
    doc.add_page_break()
    add_overview_tables(doc)
    add_login_section(doc)
    add_menu_map(doc)
    add_file_manager(doc)
    add_library_section(doc)
    add_vr_sections(doc)
    add_masterplan_and_settings(doc)
    add_json_editor(doc)
    add_workflows(doc)
    add_troubleshooting(doc)
    add_security_notes(doc)
    add_appendix(doc)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.text = "Huong dan su dung admin Bich Dong Lakeside"

    doc.save(OUT_FILE)
    return OUT_FILE


if __name__ == "__main__":
    path = build_document()
    print(path)
